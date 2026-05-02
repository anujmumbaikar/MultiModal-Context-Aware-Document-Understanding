from __future__ import annotations

import base64
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI
from tqdm import tqdm

load_dotenv()

openai_client = OpenAI()

MIN_IMAGE_BYTES = 1024
MAX_CAPTION_WORKERS = 5
PAGE_BATCH_SIZE = 50  # pages processed per batch when images/tables enabled


# ---------------------------------------------------------------------------
# Text chunking
# ---------------------------------------------------------------------------

def process_text_chunks(texts: list[dict]) -> list[dict]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    processed = []
    for item in texts:
        text = item["text"].strip()
        if not text:
            continue
        for chunk in splitter.split_text(text):
            processed.append({
                "content": chunk,
                "content_type": "text",
                "filename": item["filename"],
                "page_number": item["page_number"],
                "caption": None,
            })
    return processed


# ---------------------------------------------------------------------------
# Image captioning (parallel)
# ---------------------------------------------------------------------------

def _caption_single_image(item: dict) -> dict | None:
    b64 = item["base64"]
    ext = item.get("ext", "png")
    prompt = (
        "Generate and describe the image in detail. "
        "Directly analyze the image and provide a detailed description without any additional text. "
        "max characters should be 100-150 words."
    )
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/{ext};base64,{b64}"}},
                ],
            }],
        )
        caption = resp.choices[0].message.content
        return {
            "content": caption,
            "content_type": "image",
            "filename": item["filename"],
            "page_number": item["page_number"],
            "caption": caption,
        }
    except Exception as e:
        print(f"  Warning: image captioning failed: {e}")
        return None


def process_images_with_caption(images: list[dict]) -> list[dict]:
    if not images:
        return []
    processed = []
    with ThreadPoolExecutor(max_workers=MAX_CAPTION_WORKERS) as executor:
        futures = {executor.submit(_caption_single_image, item): item for item in images}
        for future in tqdm(as_completed(futures), total=len(images), desc="    Captioning images", unit="img", leave=False):
            result = future.result()
            if result:
                processed.append(result)
    return processed


# ---------------------------------------------------------------------------
# Table description (parallel)
# ---------------------------------------------------------------------------

def _describe_single_table(item: dict) -> dict | None:
    prompt = (
        "Generate and describe the table in detailed description of its contents, "
        "including the structure, key data points, notable trends or insights. "
        f"The table is:\n{item['markdown']}\n"
        "Directly analyze the table and provide a detailed description without any additional text. "
        "max characters should be 100-150 words."
    )
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[{"role": "user", "content": prompt}],
        )
        return {
            "content": resp.choices[0].message.content,
            "content_type": "table",
            "filename": item["filename"],
            "page_number": item["page_number"],
            "caption": None,
        }
    except Exception as e:
        print(f"  Warning: table description failed: {e}")
        return None


def process_tables_with_description(tables: list[dict]) -> list[dict]:
    if not tables:
        return []
    processed = []
    with ThreadPoolExecutor(max_workers=MAX_CAPTION_WORKERS) as executor:
        futures = {executor.submit(_describe_single_table, item): item for item in tables}
        for future in tqdm(as_completed(futures), total=len(tables), desc="    Describing tables", unit="tbl", leave=False):
            result = future.result()
            if result:
                processed.append(result)
    return processed


# ---------------------------------------------------------------------------
# PDF extraction — page-batch streaming
# ---------------------------------------------------------------------------

def _extract_pdf_streaming(file_path: str, process_images: bool, process_tables: bool) -> list[dict]:
    """
    Processes the PDF in batches of PAGE_BATCH_SIZE pages.
    For each batch: extract text/images/tables → immediately dispatch GPT-4
    calls in parallel → accumulate chunks. Memory stays bounded regardless
    of document size.
    """
    filename = os.path.basename(file_path)
    all_chunks: list[dict] = []
    seen_xrefs: set[int] = set()

    with fitz.open(file_path) as doc:
        total_pages = len(doc)
        total_batches = (total_pages + PAGE_BATCH_SIZE - 1) // PAGE_BATCH_SIZE
        print(
            f"\n[Extract] {filename} — {total_pages} pages | "
            f"images={process_images}  tables={process_tables} | "
            f"batch_size={PAGE_BATCH_SIZE} ({total_batches} batches)"
        )

        batch_bar = tqdm(
            range(0, total_pages, PAGE_BATCH_SIZE),
            total=total_batches,
            desc="[Batches]",
            unit="batch",
        )

        for batch_start in batch_bar:
            batch_end = min(batch_start + PAGE_BATCH_SIZE, total_pages)
            batch_num = batch_start // PAGE_BATCH_SIZE + 1
            batch_bar.set_postfix(pages=f"{batch_start+1}-{batch_end}")

            batch_texts: list[dict] = []
            batch_images: list[dict] = []
            batch_tables: list[dict] = []

            # --- extraction phase (sequential — fitz is not thread-safe) ---
            for page_idx in range(batch_start, batch_end):
                page = doc[page_idx]
                page_num = page_idx + 1

                for block in page.get_text("blocks"):
                    text = block[4].strip()
                    if text:
                        batch_texts.append({"text": text, "page_number": page_num, "filename": filename})

                if process_images:
                    for img_info in page.get_images(full=True):
                        xref = img_info[0]
                        if xref in seen_xrefs:
                            continue
                        seen_xrefs.add(xref)
                        img_data = doc.extract_image(xref)
                        img_bytes = img_data["image"]
                        if len(img_bytes) < MIN_IMAGE_BYTES:
                            continue
                        batch_images.append({
                            "base64": base64.b64encode(img_bytes).decode("utf-8"),
                            "ext": img_data.get("ext", "png"),
                            "page_number": page_num,
                            "filename": filename,
                        })

                if process_tables:
                    for table in page.find_tables().tables:
                        rows = table.extract()
                        if not rows:
                            continue
                        md_rows = [
                            "| " + " | ".join(str(c or "").strip() for c in row) + " |"
                            for row in rows
                        ]
                        batch_tables.append({
                            "markdown": "\n".join(md_rows),
                            "page_number": page_num,
                            "filename": filename,
                        })

            # --- processing phase ---
            text_chunks = process_text_chunks(batch_texts)
            all_chunks.extend(text_chunks)

            if batch_images:
                tqdm.write(f"  [Batch {batch_num}/{total_batches}] {len(batch_images)} images → captioning (parallel)...")
                all_chunks.extend(process_images_with_caption(batch_images))

            if batch_tables:
                tqdm.write(f"  [Batch {batch_num}/{total_batches}] {len(batch_tables)} tables → describing (parallel)...")
                all_chunks.extend(process_tables_with_description(batch_tables))

            tqdm.write(
                f"  [Batch {batch_num}/{total_batches}] pages {batch_start+1}-{batch_end} done — "
                f"{len(text_chunks)} text, {len(batch_images)} img, {len(batch_tables)} tbl chunks"
            )

    return all_chunks


# ---------------------------------------------------------------------------
# DOCX extraction (small files, no streaming needed)
# ---------------------------------------------------------------------------

def extract_from_docx(file_path: str, process_images: bool = False, process_tables: bool = False) -> dict:
    filename = os.path.basename(file_path)
    texts, images, tables = [], [], []

    document = DocxDocument(file_path)

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            texts.append({"text": text, "page_number": None, "filename": filename})

    if process_tables:
        for table in document.tables:
            md_rows = ["| " + " | ".join(cell.text.strip() for cell in row.cells) + " |"
                       for row in table.rows]
            tables.append({"markdown": "\n".join(md_rows), "page_number": None, "filename": filename})

    if process_images:
        for rel in document.part.rels.values():
            if "/image" in rel.reltype:
                img_bytes = rel.target_part.blob
                if len(img_bytes) < MIN_IMAGE_BYTES:
                    continue
                ext = rel.target_part.content_type.split("/")[-1]
                images.append({
                    "base64": base64.b64encode(img_bytes).decode("utf-8"),
                    "ext": ext,
                    "page_number": None,
                    "filename": filename,
                })

    return {"texts": texts, "images": images, "tables": tables}


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def get_all_chunks(file_path: str, process_images: bool = False, process_tables: bool = False) -> list[dict]:
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        # Streaming batch pipeline — works efficiently at any page count
        all_chunks = _extract_pdf_streaming(file_path, process_images=process_images, process_tables=process_tables)

    elif ext == ".docx":
        raw = extract_from_docx(file_path, process_images=process_images, process_tables=process_tables)
        print(f"\n[Chunking] Splitting {len(raw['texts'])} text blocks...")
        all_chunks = process_text_chunks(raw["texts"])
        all_chunks += process_images_with_caption(raw["images"])
        all_chunks += process_tables_with_description(raw["tables"])

    else:
        raise ValueError(f"Unsupported file type: {ext}")

    text_count = sum(1 for c in all_chunks if c["content_type"] == "text")
    img_count = sum(1 for c in all_chunks if c["content_type"] == "image")
    tbl_count = sum(1 for c in all_chunks if c["content_type"] == "table")
    print(f"\n[Done] {len(all_chunks)} total chunks — {text_count} text | {img_count} image | {tbl_count} table")
    return all_chunks
